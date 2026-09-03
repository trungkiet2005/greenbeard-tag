"""Generate the DOCX submission items Editorial Manager sometimes insists on.

Highlights and the Declaration of Interest are uploaded as separate items, and
some Elsevier portals reject a PDF for them. Rather than keep a hand-written
Word file that can drift from the manuscript, derive all three from the sources
that are already checked:

  highlights.docx            <- the \\item lines in highlights.tex
  declaration_of_interest.docx <- declaration_of_interest.txt
  cover_letter.docx          <- cover_letter.md, via pandoc

The highlights are re-checked against Elsevier's 85-character cap here too, so
the DOCX cannot go out over the limit even if someone edits the .tex and skips
assemble.py.

Run from paper-physicaa/, after build.py:  python make_submission_files.py
"""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

HERE = Path(__file__).resolve().parent

TITLE = ("Certification without a safety dividend: forgeable credentials in "
         "populations of autonomous agents")


def highlights_docx() -> Path:
    import docx

    src = (HERE / "highlights.tex").read_text(encoding="utf-8")
    bullets = re.findall(r"^\s*\\item\s+(.*\S)\s*$", src, re.M)
    if not 3 <= len(bullets) <= 5:
        raise SystemExit("Elsevier wants three to five highlights; found %d"
                         % len(bullets))
    over = [b for b in bullets if len(b) > 85]
    if over:
        raise SystemExit("highlights over the 85-character cap: %r" % over)

    d = docx.Document()
    d.add_heading("Highlights", level=1)
    d.add_paragraph(TITLE)
    for b in bullets:
        d.add_paragraph(b, style="List Bullet")
    out = HERE / "highlights.docx"
    d.save(out)
    return out


def declaration_docx() -> Path:
    import docx

    text = (HERE / "declaration_of_interest.txt").read_text(encoding="utf-8")
    d = docx.Document()
    for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
        one_line = " ".join(block.split())
        if block.startswith("Declaration of interests"):
            d.add_heading(one_line, level=1)
        else:
            d.add_paragraph(one_line)
    out = HERE / "declaration_of_interest.docx"
    d.save(out)
    return out


def cover_letter_docx() -> Path:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise SystemExit("pandoc not found; cover_letter.docx not generated")
    out = HERE / "cover_letter.docx"
    subprocess.run([pandoc, "cover_letter.md", "-o", out.name],
                   cwd=HERE, check=True)
    return out


def main() -> int:
    for fn in (highlights_docx, declaration_docx, cover_letter_docx):
        path = fn()
        print("wrote %-30s %6.1f kB" % (path.name, path.stat().st_size / 1024))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
