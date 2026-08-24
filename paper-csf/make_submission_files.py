"""Produce the separate editable files Chaos, Solitons & Fractals asks for at
submission: highlights, declarations and, when its private Markdown source is
present, a polished cover letter.

Highlights and declarations are separate Editorial Manager items; the cover
letter is generated only from its private authoring source.

Run from paper-csf/:  python make_submission_files.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

HERE = Path(__file__).resolve().parent

TITLE = ("Forgeable greenbeards: multistability and hollow collapse in "
         "certified agent populations")

AI_DECLARATION = (
    "During the preparation of this work the authors used Anthropic's Claude "
    "through the Claude Code command-line interface and OpenAI Codex to assist "
    "with drafting and editing prose; inspecting and revising LaTeX and "
    "bibliography files; reviewing and revising analysis and theory code, "
    "validation checks and tests; and developing deterministic plotting, build, "
    "packaging and reproducibility tooling. All numerical values reported in the "
    "manuscript were checked against or regenerated from deterministic, "
    "version-controlled scientific code, and all submitted figures were rendered "
    "deterministically from model outputs and data by that code. No generative "
    "image is included in the submission. The authors reviewed and edited all "
    "tool-assisted output, reran the analysis and test suites, reproduced the "
    "reported results, and take full responsibility for the content of the "
    "published article."
)


def _highlights_from_front_matter():
    r"""The highlights the typeset paper carries, read out of _front.tex.

    They used to be a second copy kept here by hand, and the two drifted: the
    manuscript said "Two disjoint attracting faces" where the uploaded .docx
    still said "The flow is bistable", and one bullet had lost its "at r=0".
    Editorial Manager takes the .docx and the reviewer reads the PDF, so the
    two must be the same sentence.  Reading one from the other is the only way
    that stays true.
    """
    front = (HERE / "_front.tex").read_text(encoding="utf-8")
    body = re.search(r"\\begin\{highlights\}(.*?)\\end\{highlights\}",
                     front, re.S)
    assert body, "no highlights environment in _front.tex"
    items = [line.strip()[len(r"\item"):].strip()
             for line in body.group(1).split("\n")
             if line.strip().startswith(r"\item")]
    assert items, "highlights environment is empty"
    return items


HIGHLIGHTS = _highlights_from_front_matter()

for h in HIGHLIGHTS:
    assert len(h) <= 85, f"highlight over 85 characters ({len(h)}): {h}"
assert 3 <= len(HIGHLIGHTS) <= 5, "CSF wants 3 to 5 highlights"

# Keep the optional standalone LaTeX page on the same source as the editable
# Word upload and the manuscript front matter.
latex_items = "\n".join(r"\item " + h for h in HIGHLIGHTS)
(HERE / "highlights.tex").write_text(
    "%% Generated from _front.tex by make_submission_files.py.\n"
    "\\documentclass[12pt]{article}\n"
    "\\usepackage[T1]{fontenc}\n"
    "\\usepackage{lmodern}\n"
    "\\usepackage[margin=1in]{geometry}\n"
    "\\pagestyle{empty}\n"
    "\\begin{document}\n"
    "\\section*{Highlights}\n"
    "\\begin{itemize}\n"
    f"{latex_items}\n"
    "\\end{itemize}\n"
    "\\end{document}\n",
    encoding="utf-8",
)


def new_doc():
    doc = Document()
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        # The default template uses bestFit but omits the percentage required
        # by the ECMA-376 schema used by the submission validator.
        zoom.set(qn("w:percent"), "100")
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


# ------------------------------------------------------------- highlights
doc = new_doc()
doc.add_heading("Highlights", level=1)
p = doc.add_paragraph()
p.add_run(TITLE).italic = True
for h in HIGHLIGHTS:
    doc.add_paragraph(h, style="List Bullet")
doc.save(HERE / "highlights.docx")

# --------------------------------------------- declaration of interests
doc = new_doc()
doc.add_heading("Declaration of competing interest", level=1)
p = doc.add_paragraph()
p.add_run("Manuscript title: ").bold = True
p.add_run(TITLE)
doc.add_paragraph(
    "The authors declare that they have no known competing financial interests "
    "or personal relationships that could have appeared to influence the work "
    "reported in this paper."
)
doc.add_heading("Funding", level=2)
doc.add_paragraph(
    "This research did not receive any specific grant from funding agencies in "
    "the public, commercial, or not-for-profit sectors."
)
doc.add_heading(
    "Declaration of generative AI and AI-assisted technologies in the "
    "manuscript preparation process",
    level=2,
)
doc.add_paragraph(AI_DECLARATION)
doc.save(HERE / "declaration_of_interest.docx")


def _add_markdown_runs(paragraph, text):
    """Render the small inline-Markdown subset used by the private letter."""
    for token in re.split(r"(\*\*.*?\*\*|\*.*?\*|<https?://[^>]+>)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("<http") and token.endswith(">"):
            paragraph.add_run(token[1:-1])
        else:
            paragraph.add_run(token)


def _markdown_blocks(text):
    blocks = []
    current = []
    for raw_line in text.splitlines():
        if not raw_line.strip():
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(raw_line)
    if current:
        blocks.append(current)
    return blocks


def write_cover_letter():
    """Generate the private upload copy without publishing its source."""
    source = HERE / "cover_letter.md"
    if not source.is_file():
        print("skipped cover_letter.docx (private cover_letter.md is absent)")
        return

    blocks = _markdown_blocks(source.read_text(encoding="utf-8"))
    assert blocks and blocks[0] == ["# Cover letter"], "unexpected cover heading"

    cover = new_doc()
    section = cover.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    normal = cover.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for index, lines in enumerate(blocks[1:]):
        segments = []
        current = ""
        for line in lines:
            hard_break = line.endswith("  ")
            cleaned = line.rstrip()
            current = f"{current} {cleaned}".strip()
            if hard_break:
                segments.append(current)
                current = ""
        if current:
            segments.append(current)

        paragraph = cover.add_paragraph()
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif len(lines) > 1 and not any(line.endswith("  ") for line in lines):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for segment_index, segment in enumerate(segments):
            if segment_index:
                paragraph.add_run().add_break()
            _add_markdown_runs(paragraph, segment)

    cover.save(HERE / "cover_letter.docx")
    print("wrote cover_letter.docx from private cover_letter.md")


write_cover_letter()
print("wrote highlights.tex, highlights.docx and declaration_of_interest.docx")
